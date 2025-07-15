"""
Word export functionality for the Document Creator toolkit.

This module provides functions to export normalized data to Word document format,
supporting both direct document generation and template-based document creation.
"""

import tempfile
import os
import requests
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
import logging
from datetime import datetime
import json

# Optional imports for Word document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docxtpl import DocxTemplate
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from .utils import (
    sanitize_filename,
    ensure_directory_exists,
    DocumentCreatorError,
    get_available_filename
)


class WordExportError(DocumentCreatorError):
    """Exception for Word export related errors."""
    pass


def export_to_word(
    data_list: List[Dict[str, Any]],
    output_directory: str,
    filename_key: Optional[str] = None,
    template_url: Optional[str] = None,
    template_path: Optional[str] = None,
    use_template: bool = False,
    document_title: Optional[str] = None,
    document_author: Optional[str] = None
) -> List[Path]:
    """
    Export a list of dictionaries to Word document format.
    
    Args:
        data_list: List of dictionaries to export
        output_directory: Directory to save Word files
        filename_key: Key to use for filename generation
        template_url: URL to Word template for template-based export
        template_path: Path to local Word template
        use_template: Whether to use template-based export
        document_title: Title for generated documents
        document_author: Author for generated documents
        
    Returns:
        List of Path objects for created files
        
    Raises:
        WordExportError: If export fails
    """
    try:
        # Validate requirements
        if not DOCX_AVAILABLE:
            raise WordExportError("Word export requires python-docx and docxtpl packages")
        
        # Ensure output directory exists
        output_path = ensure_directory_exists(output_directory)
        
        # Validate input data
        if not isinstance(data_list, list):
            raise WordExportError("Data must be a list of dictionaries")
        
        if not data_list:
            raise WordExportError("No data provided for export")
        
        created_files = []
        
        # Choose export method
        if use_template:
            created_files = export_word_with_template(
                data_list,
                output_path,
                filename_key,
                template_url,
                template_path
            )
        else:
            created_files = export_word_direct(
                data_list,
                output_path,
                filename_key,
                document_title,
                document_author
            )
        
        return created_files
        
    except Exception as e:
        raise WordExportError(f"Failed to export to Word: {str(e)}")


def export_word_direct(
    data_list: List[Dict[str, Any]],
    output_path: Path,
    filename_key: Optional[str],
    document_title: Optional[str],
    document_author: Optional[str]
) -> List[Path]:
    """
    Export data to Word using direct generation (python-docx).
    
    Args:
        data_list: List of dictionaries to export
        output_path: Directory to save Word documents
        filename_key: Key to use for filename generation
        document_title: Title for generated documents
        document_author: Author for generated documents
        
    Returns:
        List of Path objects for created files
    """
    created_files = []
    
    for index, data_object in enumerate(data_list):
        if not isinstance(data_object, dict):
            logging.warning(f"Skipping non-dictionary item at index {index}")
            continue
        
        # Generate filename
        filename = generate_word_filename(data_object, filename_key, index)
        file_path = get_available_filename(output_path / filename, "docx")
        
        # Create Word document
        create_word_document(
            data_object,
            file_path,
            document_title,
            document_author
        )
        
        created_files.append(file_path)
        logging.info(f"✅ Created: {file_path.name}")
    
    return created_files


def export_word_with_template(
    data_list: List[Dict[str, Any]],
    output_path: Path,
    filename_key: Optional[str],
    template_url: Optional[str],
    template_path: Optional[str]
) -> List[Path]:
    """
    Export data to Word using template conversion.
    
    Args:
        data_list: List of dictionaries to export
        output_path: Directory to save Word documents
        filename_key: Key to use for filename generation
        template_url: URL to Word template
        template_path: Path to local Word template
        
    Returns:
        List of Path objects for created files
    """
    created_files = []
    
    # Get template
    if template_url:
        template_data = download_template(template_url)
    elif template_path:
        with open(template_path, 'rb') as f:
            template_data = f.read()
    else:
        raise WordExportError("Either template_url or template_path must be provided")
    
    for index, data_object in enumerate(data_list):
        if not isinstance(data_object, dict):
            logging.warning(f"Skipping non-dictionary item at index {index}")
            continue
        
        # Generate filename
        filename = generate_word_filename(data_object, filename_key, index)
        file_path = get_available_filename(output_path / filename, "docx")
        
        # Create Word document from template
        word_data = render_template_to_word(template_data, data_object)
        
        # Write Word file
        with open(file_path, 'wb') as f:
            f.write(word_data)
        
        created_files.append(file_path)
        logging.info(f"✅ Created: {file_path.name}")
    
    return created_files


def generate_word_filename(
    data_object: Dict[str, Any],
    filename_key: Optional[str],
    index: int
) -> str:
    """
    Generate a filename for a Word file based on data object.
    
    Args:
        data_object: Dictionary containing data
        filename_key: Key to use for filename generation
        index: Index of the object in the list
        
    Returns:
        Sanitized filename without extension
    """
    base_filename = None
    
    # Try to get filename from specified key
    if filename_key:
        base_filename = data_object.get(filename_key)
    
    # Try common fallback keys
    if not base_filename:
        fallback_keys = ["name", "title", "filename", "id"]
        for key in fallback_keys:
            if key in data_object and data_object[key]:
                base_filename = data_object[key]
                break
    
    # Use timestamp + index as final fallback
    if not base_filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"document_{timestamp}_{index + 1}"
    
    return sanitize_filename(str(base_filename))


def create_word_document(
    data_object: Dict[str, Any],
    file_path: Path,
    title: Optional[str] = None,
    author: Optional[str] = None
) -> None:
    """
    Create a Word document using python-docx from a data object.
    
    Args:
        data_object: Dictionary containing data to include in document
        file_path: Path where document should be saved
        title: Title for the document
        author: Author for the document
    """
    doc = Document()
    
    # Set document properties
    if title:
        doc.core_properties.title = title
    if author:
        doc.core_properties.author = author
    
    # Add title
    if title:
        title_paragraph = doc.add_heading(title, level=1)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content sections
    for key, value in data_object.items():
        # Add section heading
        section_title = key.replace('_', ' ').title()
        doc.add_heading(section_title, level=2)
        
        # Add content based on value type
        if isinstance(value, dict):
            # Add dictionary as table
            add_dict_as_table(doc, value)
        elif isinstance(value, list):
            # Add list as bulleted list
            add_list_as_bullets(doc, value)
        else:
            # Add simple value as paragraph
            doc.add_paragraph(str(value))
        
        # Add some spacing
        doc.add_paragraph()
    
    # Save document
    doc.save(str(file_path))


def add_dict_as_table(doc: Document, data_dict: Dict[str, Any]) -> None:
    """
    Add a dictionary as a table to a Word document.
    
    Args:
        doc: Word document object
        data_dict: Dictionary to add as table
    """
    if not data_dict:
        doc.add_paragraph("No data available")
        return
    
    # Create table
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # Add header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Key'
    header_cells[1].text = 'Value'
    
    # Add data rows
    for key, value in data_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        if isinstance(value, (dict, list)):
            row_cells[1].text = json.dumps(value, indent=2, ensure_ascii=False)
        else:
            row_cells[1].text = str(value)


def add_list_as_bullets(doc: Document, data_list: List[Any]) -> None:
    """
    Add a list as bulleted list to a Word document.
    
    Args:
        doc: Word document object
        data_list: List to add as bullets
    """
    if not data_list:
        doc.add_paragraph("No items available")
        return
    
    for item in data_list:
        if isinstance(item, (dict, list)):
            # Format complex items as JSON
            item_text = json.dumps(item, indent=2, ensure_ascii=False)
        else:
            item_text = str(item)
        
        doc.add_paragraph(item_text, style='List Bullet')


def download_template(template_url: str) -> bytes:
    """
    Download a Word template from URL.
    
    Args:
        template_url: URL to the Word template
        
    Returns:
        Template data as bytes
        
    Raises:
        WordExportError: If template cannot be downloaded
    """
    try:
        response = requests.get(template_url)
        response.raise_for_status()
        return response.content
    except requests.RequestException as e:
        raise WordExportError(f"Failed to download template: {str(e)}")


def render_template_to_word(
    template_data: bytes,
    data_object: Dict[str, Any],
    opening_delimiter: str = "{{",
    closing_delimiter: str = "}}",
    use_docxtpl: bool = True
) -> bytes:
    """
    Render a Word template with data.
    
    Args:
        template_data: Word template data as bytes
        data_object: Dictionary containing data for template
        opening_delimiter: Opening delimiter for template variables
        closing_delimiter: Closing delimiter for template variables
        use_docxtpl: Whether to use docxtpl for advanced templating
        
    Returns:
        Word document data as bytes
        
    Raises:
        WordExportError: If template rendering fails
    """
    try:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Save template
            template_path = temp_path / "template.docx"
            with open(template_path, 'wb') as f:
                f.write(template_data)
            
            # Render template
            if use_docxtpl:
                doc = DocxTemplate(str(template_path))
                doc.render(data_object)
                rendered_path = temp_path / "rendered.docx"
                doc.save(str(rendered_path))
            else:
                # Simple variable replacement
                doc = Document(str(template_path))
                for paragraph in doc.paragraphs:
                    for key, value in data_object.items():
                        placeholder = f"{opening_delimiter}{key}{closing_delimiter}"
                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
                
                rendered_path = temp_path / "rendered.docx"
                doc.save(str(rendered_path))
            
            # Read rendered document data
            with open(rendered_path, 'rb') as f:
                return f.read()
                
    except Exception as e:
        raise WordExportError(f"Failed to render template to Word: {str(e)}")


def check_word_requirements() -> Dict[str, bool]:
    """
    Check which Word export features are available.
    
    Returns:
        Dictionary indicating which features are available
    """
    return {
        'word_export': DOCX_AVAILABLE,
        'python_docx': DOCX_AVAILABLE,
        'template_support': DOCX_AVAILABLE
    }


def get_missing_requirements() -> List[str]:
    """
    Get list of missing packages for Word export.
    
    Returns:
        List of missing package names
    """
    missing = []
    
    if not DOCX_AVAILABLE:
        missing.extend(['python-docx', 'docxtpl'])
    
    return missing
