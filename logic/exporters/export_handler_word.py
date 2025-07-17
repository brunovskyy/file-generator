"""
Word document exporter for DocGenius.

This module provides Word document export functionality using the BaseExporter
architecture with support for direct Word generation and template-based
Word document creation.
"""

import logging
from typing import List, Dict, Any, Optional, Union
from pathlib import Path
from datetime import datetime

from ..models.base_models import ValidationResult
from ..models.data_structures import DataObject, DataCollection
from ..models.document_config import DocumentConfig, WordSettings
from .export_handler_base import BaseExporter, ExportResult, ExportContext
from .template_processor import TemplateProcessor


class WordExporter(BaseExporter):
    """Word document exporter with template support."""
    
    def __init__(self):
        super().__init__()
        self.format_name = "Word"
        self.file_extension = ".docx"
        self.template_processor = TemplateProcessor()
        
        # Check Word dependencies
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if Word generation dependencies are available."""
        self.dependencies_available = True
        self.missing_dependencies = []
        
        try:
            import docx
            self.logger.debug("python-docx available for Word generation")
        except ImportError:
            self.dependencies_available = False
            self.missing_dependencies.append("python-docx")
            self.logger.warning("python-docx not available - Word generation disabled")
        
        # Check for template support
        try:
            import docxtpl
            self.template_support = True
            self.logger.debug("docxtpl available for template processing")
        except ImportError:
            self.template_support = False
            self.missing_dependencies.append("docxtpl")
            self.logger.debug("docxtpl not available - template support disabled")
    
    def validate_config(self, config: DocumentConfig) -> ValidationResult:
        """
        Validate Word-specific configuration.
        
        Args:
            config: Document configuration to validate
            
        Returns:
            ValidationResult with validation details
        """
        errors = []
        warnings = []
        
        # Check base configuration
        base_result = super().validate_config(config)
        if not base_result.is_valid:
            errors.extend(base_result.errors)
        if base_result.warnings:
            warnings.extend(base_result.warnings)
        
        # Check Word-specific settings
        if hasattr(config, 'word_settings') and config.word_settings:
            word_settings = config.word_settings
            
            # Validate template file if specified
            if config.template_file:
                template_path = Path(config.template_file)
                if not template_path.exists():
                    errors.append(f"Template file not found: {config.template_file}")
                elif template_path.suffix.lower() not in ['.docx', '.dotx']:
                    warnings.append(f"Template file should be .docx or .dotx format: {template_path.suffix}")
            
            # Validate style mappings
            if hasattr(word_settings, 'style_mapping') and word_settings.style_mapping:
                if not isinstance(word_settings.style_mapping, dict):
                    errors.append("Word style mapping must be a dictionary")
        
        # Check template support if template is specified
        if config.template_file and not self.template_support:
            warnings.append("Template specified but docxtpl not available - using basic generation")
        
        return ValidationResult(
            is_valid=len(errors) == 0,
            errors=errors,
            warnings=warnings
        )
    
    def export_single(self, data_object: DataObject, context: ExportContext) -> ExportResult:
        """
        Export a single data object to Word document.
        
        Args:
            data_object: Data object to export
            context: Export context with configuration
            
        Returns:
            ExportResult with export details
        """
        try:
            if not self.dependencies_available:
                return ExportResult(
                    success=False,
                    error_message=f"Word dependencies not available: {', '.join(self.missing_dependencies)}",
                    exported_files=[]
                )
            
            # Generate filename
            filename = self._generate_filename(data_object, context.config)
            output_path = context.output_directory / f"{filename}.docx"
            
            # Ensure output directory exists
            output_path.parent.mkdir(parents=True, exist_ok=True)
            
            # Choose export method based on configuration
            if context.config.template_file and self.template_support:
                success = self._export_with_template(data_object, output_path, context)
            else:
                success = self._export_direct(data_object, output_path, context)
            
            if success:
                return ExportResult(
                    success=True,
                    exported_files=[output_path]
                )
            else:
                return ExportResult(
                    success=False,
                    error_message="Word document generation failed",
                    exported_files=[]
                )
        
        except Exception as e:
            self.logger.error(f"Word export failed for {data_object.get_display_name()}: {str(e)}")
            return ExportResult(
                success=False,
                error_message=str(e),
                exported_files=[]
            )
    
    def _export_direct(self, data_object: DataObject, output_path: Path, 
                      context: ExportContext) -> bool:
        """
        Generate Word document directly using python-docx.
        
        Args:
            data_object: Data object to export
            output_path: Output file path
            context: Export context
            
        Returns:
            True if successful
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # Create new document
            doc = Document()
            
            # Get Word settings
            word_settings = getattr(context.config, 'word_settings', WordSettings())
            
            # Configure document styles
            self._configure_document_styles(doc, word_settings)
            
            # Add title
            title = data_object.get_display_name()
            title_paragraph = doc.add_heading(title, level=1)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata if enabled
            if context.config.include_metadata:
                self._add_metadata_section(doc, context)
            
            # Add main content
            self._add_content_section(doc, data_object, word_settings)
            
            # Save document
            doc.save(str(output_path))
            
            self.logger.info(f"Generated Word document: {output_path}")
            return True
        
        except Exception as e:
            self.logger.error(f"Direct Word generation failed: {str(e)}")
            return False
    
    def _export_with_template(self, data_object: DataObject, output_path: Path,
                             context: ExportContext) -> bool:
        """
        Generate Word document using a template.
        
        Args:
            data_object: Data object to export
            output_path: Output file path
            context: Export context
            
        Returns:
            True if successful
        """
        try:
            from docxtpl import DocxTemplate
            
            template_path = Path(context.config.template_file)
            
            if not template_path.exists():
                self.logger.error(f"Template file not found: {template_path}")
                return False
            
            # Load template
            doc = DocxTemplate(str(template_path))
            
            # Prepare context data
            template_context = data_object.get_flattened_data()
            
            # Add metadata
            if context.config.include_metadata:
                template_context.update({
                    'generated_date': datetime.now().strftime('%Y-%m-%d'),
                    'generated_time': datetime.now().strftime('%H:%M:%S'),
                    'generated_datetime': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                    'source': getattr(context.config, 'source_info', 'Unknown'),
                    'format': 'Word Document',
                    'exporter': 'DocGenius Word Exporter'
                })
            
            # Add special formatting helpers
            template_context.update({
                'current_date': datetime.now().strftime('%Y-%m-%d'),
                'current_time': datetime.now().strftime('%H:%M:%S'),
                'document_title': data_object.get_display_name()
            })
            
            # Handle nested data for template
            self._prepare_template_context(template_context, data_object)
            
            # Render template
            doc.render(template_context)
            
            # Save document
            doc.save(str(output_path))
            
            self.logger.info(f"Generated Word document from template: {output_path}")
            return True
        
        except Exception as e:
            self.logger.error(f"Template-based Word generation failed: {str(e)}")
            return False
    
    def _configure_document_styles(self, doc, word_settings: WordSettings) -> None:
        """Configure document-wide styles."""
        try:
            from docx.shared import Pt
            
            # Configure default styles
            styles = doc.styles
            
            # Normal style
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = getattr(word_settings, 'font_family', 'Calibri')
            normal_font.size = Pt(getattr(word_settings, 'font_size', 11))
            
            # Heading styles
            heading1_style = styles['Heading 1']
            heading1_font = heading1_style.font
            heading1_font.name = getattr(word_settings, 'heading_font_family', 'Calibri')
            heading1_font.size = Pt(getattr(word_settings, 'heading_font_size', 16))
            
        except Exception as e:
            self.logger.warning(f"Failed to configure document styles: {str(e)}")
    
    def _add_metadata_section(self, doc, context: ExportContext) -> None:
        """Add metadata section to document."""
        try:
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Add metadata heading
            doc.add_heading('Document Information', level=2)
            
            # Create metadata table
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Add metadata rows
            metadata_items = [
                ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Format', 'Word Document'),
                ('Source', getattr(context.config, 'source_info', 'Unknown')),
                ('Exporter', 'DocGenius Word Exporter')
            ]
            
            for key, value in metadata_items:
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = str(value)
                
                # Style the key cell
                key_paragraph = row.cells[0].paragraphs[0]
                key_run = key_paragraph.runs[0]
                key_run.bold = True
            
            # Add space after metadata
            doc.add_paragraph()
        
        except Exception as e:
            self.logger.warning(f"Failed to add metadata section: {str(e)}")
    
    def _add_content_section(self, doc, data_object: DataObject, 
                           word_settings: WordSettings) -> None:
        """Add main content section to document."""
        try:
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches
            
            # Add content heading
            doc.add_heading('Content', level=2)
            
            # Get flattened data
            flat_data = data_object.get_flattened_data()
            
            # Create content table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Header row
            header_row = table.rows[0]
            header_row.cells[0].text = 'Field'
            header_row.cells[1].text = 'Value'
            
            # Style header row
            for cell in header_row.cells:
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0]
                run.bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add data rows
            for key, value in flat_data.items():
                row = table.add_row()
                row.cells[0].text = str(key)
                
                # Handle complex values
                if isinstance(value, (dict, list)):
                    import json
                    value_str = json.dumps(value, indent=2, default=str)
                else:
                    value_str = str(value) if value is not None else ''
                
                row.cells[1].text = value_str
                
                # Style the field name
                field_paragraph = row.cells[0].paragraphs[0]
                if field_paragraph.runs:
                    field_run = field_paragraph.runs[0]
                    field_run.bold = True
            
            # Adjust column widths
            for row in table.rows:
                row.cells[0].width = Inches(2.0)
                row.cells[1].width = Inches(4.0)
        
        except Exception as e:
            self.logger.warning(f"Failed to add content section: {str(e)}")
    
    def _prepare_template_context(self, template_context: Dict[str, Any], 
                                 data_object: DataObject) -> None:
        """Prepare template context with additional formatting helpers."""
        try:
            # Add original nested data
            template_context['data'] = data_object.data
            
            # Add formatted lists for common patterns
            if isinstance(data_object.data, dict):
                # Create formatted field lists
                template_context['field_list'] = [
                    {'name': k, 'value': str(v) if v is not None else ''}
                    for k, v in data_object.data.items()
                ]
                
                # Create key-value pairs for easy iteration
                template_context['fields'] = list(data_object.data.items())
        
        except Exception as e:
            self.logger.warning(f"Failed to prepare template context: {str(e)}")
    
    def get_export_capabilities(self) -> Dict[str, Any]:
        """Get Word export capabilities."""
        capabilities = super().get_export_capabilities()
        
        capabilities.update({
            "template_support": self.template_support,
            "direct_generation": True,
            "style_support": True,
            "table_support": True,
            "image_support": True,
            "header_footer_support": True,
            "page_formatting": True,
            "dependencies": {
                "required": ["python-docx"],
                "optional": ["docxtpl"],
                "available": self.dependencies_available,
                "template_support": self.template_support
            }
        })
        
        return capabilities


# Backward compatibility functions
def export_to_word(data_list: List[Dict[str, Any]], 
                  output_directory: Union[str, Path],
                  filename_key: Optional[str] = None,
                  template_file: Optional[str] = None,
                  **kwargs) -> List[Path]:
    """
    Backward compatibility function for Word export.
    
    Args:
        data_list: List of data dictionaries
        output_directory: Output directory path
        filename_key: Key to use for filename generation
        template_file: Optional template file path
        **kwargs: Additional configuration options
        
    Returns:
        List of generated file paths
    """
    try:
        # Convert to new data structures
        data_objects = [DataObject(data) for data in data_list]
        data_collection = DataCollection(data_objects)
        
        # Create configuration
        config = DocumentConfig(
            output_format="word",
            filename_key=filename_key,
            template_file=template_file,
            include_metadata=kwargs.get('include_metadata', True)
        )
        
        # Set Word-specific settings if provided
        if 'word_settings' in kwargs:
            config.word_settings = kwargs['word_settings']
        elif any(k.startswith('word_') for k in kwargs):
            # Extract Word settings from kwargs
            word_settings = WordSettings()
            if 'word_font_family' in kwargs:
                word_settings.font_family = kwargs['word_font_family']
            if 'word_font_size' in kwargs:
                word_settings.font_size = kwargs['word_font_size']
            if 'word_style_mapping' in kwargs:
                word_settings.style_mapping = kwargs['word_style_mapping']
            config.word_settings = word_settings
        
        # Create exporter and export
        exporter = WordExporter()
        results = exporter.export_collection(data_collection, config, Path(output_directory))
        
        # Extract successful file paths
        exported_files = []
        for result in results:
            if result.success:
                exported_files.extend(result.exported_files)
        
        return exported_files
    
    except Exception as e:
        logging.getLogger(__name__).error(f"Word export failed: {str(e)}")
        return []


def check_word_requirements() -> bool:
    """Check if Word export requirements are met."""
    exporter = WordExporter()
    return exporter.dependencies_available


def get_missing_requirements() -> List[str]:
    """Get list of missing Word export requirements."""
    exporter = WordExporter()
    return exporter.missing_dependencies


def check_template_support() -> bool:
    """Check if Word template support is available."""
    exporter = WordExporter()
    return exporter.template_support
